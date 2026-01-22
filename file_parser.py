"""
Enhanced File Parser Module
Combines robust parsing with advanced validation
Features from both implementations
"""

import io
import re
from PyPDF2 import PdfReader
from docx import Document
import streamlit as st

def clean_text(text):
    """
    Advanced text cleaning and normalization
    
    Args:
        text (str): Raw text to clean
    
    Returns:
        str: Cleaned and normalized text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove very long sequences of special characters
    text = re.sub(r'[^\w\s.,;:()\-\/]+', '', text)
    
    # Remove carriage returns
    text = text.replace('\r', '')
    
    # Split into lines and remove duplicates while preserving order
    lines = text.split('\n')
    seen = set()
    unique_lines = []
    for line in lines:
        line = line.strip()
        if line and line not in seen:
            seen.add(line)
            unique_lines.append(line)
    
    return '\n'.join(unique_lines)

def parse_pdf(file):
    """
    Extract text from PDF with advanced validation
    
    Args:
        file: Uploaded PDF file object
    
    Returns:
        str: Extracted text or None if failed
    """
    try:
        pdf_reader = PdfReader(file)
        
        # Check if PDF is encrypted (from your implementation)
        if pdf_reader.is_encrypted:
            st.error(
                "🔒 **This PDF is encrypted or password-protected.**\n\n"
                "Unable to extract text. Please provide an unprotected version."
            )
            return None
        
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                content = page.extract_text()
                if content:
                    text += content + "\n"
            except Exception as e:
                st.warning(f"⚠️ Could not extract text from page {page_num + 1}: {str(e)}")
                continue
        
        if not text.strip():
            st.warning(
                "📄 **No text content found in PDF.**\n\n"
                "This might be a scanned document. Please use a searchable PDF or paste text manually."
            )
            return None
        
        return clean_text(text)
        
    except Exception as e:
        st.error(f"❌ **Error reading PDF:** {str(e)}")
        return None

def parse_docx(file):
    """
    Extract text from DOCX file with error handling
    
    Args:
        file: Uploaded DOCX file object
    
    Returns:
        str: Extracted text or None if failed
    """
    try:
        doc = Document(file)
        text = ""
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract from tables (bonus feature)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            st.warning("📄 **No text content found in DOCX file.**")
            return None
        
        return clean_text(text)
        
    except Exception as e:
        st.error(
            f"❌ **Unable to read DOCX file.**\n\n"
            f"Error: {str(e)}\n\n"
            "The file may be corrupted, encrypted, or in an unsupported format."
        )
        return None

def parse_txt(file):
    """
    Extract text from TXT file with encoding detection
    
    Args:
        file: Uploaded TXT file object
    
    Returns:
        str: Extracted text or None if failed
    """
    try:
        # Try UTF-8 first
        text = file.read().decode('utf-8')
        return clean_text(text)
    except UnicodeDecodeError:
        # Fallback to latin-1
        try:
            file.seek(0)
            text = file.read().decode('latin-1')
            st.info("ℹ️ File decoded using latin-1 encoding")
            return clean_text(text)
        except Exception as e:
            st.error(f"❌ **Unable to decode text file:** {str(e)}")
            return None
    except Exception as e:
        st.error(f"❌ **Error reading TXT file:** {str(e)}")
        return None

def parse_file(file):
    """
    Main file parsing function with comprehensive validation
    Supports PDF, DOCX, and TXT files
    
    Args:
        file: Uploaded file object from Streamlit
    
    Returns:
        str: Extracted text or None if parsing fails
    """
    if file is None:
        return None
    
    file_name = file.name.lower()
    file_type = file.type
    
    # ===================================
    # IMAGE FILE DETECTION (From your code)
    # ===================================
    if file_name.endswith(('.jpg', '.jpeg', '.png', '.webp', '.gif', '.bmp', '.tiff')):
        st.error(
            "🖼️ **Image files are not supported for text parsing.**\n\n"
            "📄 Please upload a text-based document:\n"
            "- PDF (searchable, not scanned)\n"
            "- DOCX (Microsoft Word)\n"
            "- TXT (Plain text)\n\n"
            "💡 **Tip:** If you have a scanned document, consider using OCR software first."
        )
        return None
    
    # Reset file pointer
    file.seek(0)
    
    # ===================================
    # FILE TYPE DETECTION & PARSING
    # ===================================
    
    # PDF Files
    if file_type == "application/pdf" or file_name.endswith('.pdf'):
        return parse_pdf(file)
    
    # DOCX Files
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.endswith('.docx'):
        return parse_docx(file)
    
    # TXT Files
    elif file_type == "text/plain" or file_name.endswith('.txt'):
        return parse_txt(file)
    
    # DOC Files (Old Word format - not supported)
    elif file_name.endswith('.doc'):
        st.error(
            "⚠️ **Legacy .DOC format is not supported.**\n\n"
            "Please convert to .DOCX format using Microsoft Word or Google Docs."
        )
        return None
    
    # RTF Files
    elif file_name.endswith('.rtf'):
        st.error(
            "⚠️ **RTF format is not supported.**\n\n"
            "Please convert to PDF or DOCX format."
        )
        return None
    
    # Unsupported Format
    else:
        st.error(
            f"❌ **Unsupported file format: `{file_name}`**\n\n"
            "**Supported formats:**\n"
            "- PDF (.pdf)\n"
            "- Word Document (.docx)\n"
            "- Plain Text (.txt)\n\n"
            "**Not supported:**\n"
            "- Images (.jpg, .png, etc.)\n"
            "- Legacy Word (.doc)\n"
            "- Rich Text (.rtf)\n"
            "- Scanned PDFs"
        )
        return None

def validate_file_size(file, max_size_mb=10):
    """
    Validate file size before processing
    
    Args:
        file: Uploaded file object
        max_size_mb (int): Maximum allowed file size in MB
    
    Returns:
        bool: True if valid, False otherwise
    """
    if file is None:
        return False
    
    file_size_mb = file.size / (1024 * 1024)
    
    if file_size_mb > max_size_mb:
        st.error(
            f"📦 **File too large: {file_size_mb:.2f} MB**\n\n"
            f"Maximum allowed size: {max_size_mb} MB\n\n"
            "💡 **Tip:** Try compressing the PDF or removing unnecessary images."
        )
        return False
    
    return True

def extract_metadata(file):
    """
    Extract metadata from uploaded file
    
    Args:
        file: Uploaded file object
    
    Returns:
        dict: File metadata
    """
    metadata = {
        'filename': file.name,
        'size_bytes': file.size,
        'size_mb': round(file.size / (1024 * 1024), 2),
        'type': file.type
    }
    
    return metadata

# ===================================
# ADVANCED FEATURES
# ===================================

def detect_language(text):
    """
    Simple language detection (can be enhanced with langdetect library)
    
    Args:
        text (str): Input text
    
    Returns:
        str: Detected language
    """
    # Basic detection - check for common English words
    english_indicators = ['the', 'and', 'is', 'in', 'to', 'of', 'a']
    text_lower = text.lower()
    
    english_count = sum(1 for word in english_indicators if word in text_lower)
    
    if english_count >= 3:
        return 'English'
    else:
        return 'Unknown'

def count_words(text):
    """
    Count words in text
    
    Args:
        text (str): Input text
    
    Returns:
        int: Word count
    """
    if not text:
        return 0
    return len(text.split())

def count_sentences(text):
    """
    Count sentences in text
    
    Args:
        text (str): Input text
    
    Returns:
        int: Sentence count
    """
    if not text:
        return 0
    return len(re.split(r'[.!?]+', text))

def get_text_statistics(text):
    """
    Get comprehensive text statistics
    
    Args:
        text (str): Input text
    
    Returns:
        dict: Text statistics
    """
    if not text:
        return {
            'characters': 0,
            'words': 0,
            'sentences': 0,
            'language': 'Unknown'
        }
    
    return {
        'characters': len(text),
        'words': count_words(text),
        'sentences': count_sentences(text),
        'language': detect_language(text)
    }